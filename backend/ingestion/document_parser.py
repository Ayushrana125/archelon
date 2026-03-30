"""
document_parser.py — Extracts raw elements from PDF and DOCX files.
Each element has a type (title, h1, h2, h3, body, list_item, code, table, caption),
text content, and page number.
"""

import os
from collections import Counter
from dataclasses import dataclass


# ─────────────────────────────────────────────────────────────────────────────
# ELEMENT SCHEMA
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Element:
    type: str   # title | h1 | h2 | h3 | body | list_item | code | table | caption
    text: str
    page: int = 0


# ─────────────────────────────────────────────────────────────────────────────
# FORMAT DETECTION
# ─────────────────────────────────────────────────────────────────────────────

def detect_format(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx": return "docx"
    if ext == ".pdf":  return "pdf"
    with open(path, "rb") as f:
        header = f.read(8)
    if header[:4] == b"%PDF":        return "pdf"
    if header[:4] == b"PK\x03\x04": return "docx"
    raise ValueError(f"Unsupported format: {ext} — only .docx and .pdf supported")


def extract_title(elements: list) -> str:
    for el in elements[:10]:
        if el.type in ("title", "h1"):
            return el.text
    for el in elements:
        if el.text.strip():
            return el.text.strip()
    return "Untitled Document"


# ─────────────────────────────────────────────────────────────────────────────
# DOCX ENGINE
# ─────────────────────────────────────────────────────────────────────────────

def _size_pt(runs):
    sizes = [r.font.size for r in runs if r.font.size]
    return round(max(sizes) / 12700) if sizes else None

def _all_bold(runs):
    tr = [r for r in runs if r.text.strip()]
    return bool(tr) and all(r.bold for r in tr)

def _all_italic(runs):
    tr = [r for r in runs if r.text.strip()]
    return bool(tr) and all(r.italic for r in tr)

def _build_size_map(raw_elements: list) -> dict:
    size_counts = Counter()
    total = 0
    for el in raw_elements:
        sz = el.get("size_pt")
        if sz:
            size_counts[sz] += 1
            total += 1
    if not size_counts:
        return {}
    body_threshold = max(total * 0.25, 2)
    body_sizes     = {sz for sz, cnt in size_counts.items() if cnt >= body_threshold}
    heading_sizes  = sorted([sz for sz in size_counts if sz not in body_sizes], reverse=True)
    size_map = {sz: "body" for sz in body_sizes}
    for i, sz in enumerate(heading_sizes):
        size_map[sz] = ["title", "h1", "h2", "h3"][min(i, 3)]
    return size_map

STYLE_ROLE = {
    "heading 1": "h1", "heading 2": "h2", "heading 3": "h3", "heading 4": "h3",
    "title": "title", "subtitle": "caption",
    "list paragraph": "list_item", "list bullet": "list_item", "list number": "list_item",
    "caption": "caption", "intense quote": "body", "quote": "body",
}

def extract_docx(path: str) -> list:
    from docx import Document
    doc      = Document(path)
    body     = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map  = {t._element: t for t in doc.tables}

    raw = []
    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            para = para_map.get(child)
            if para and para.text.strip():
                raw.append({"size_pt": _size_pt(para.runs)})
        elif tag == 'tbl':
            tbl = tbl_map.get(child)
            if tbl:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                raw.append({"size_pt": _size_pt(p.runs)})
                                break

    size_map = _build_size_map(raw)
    elements = []

    def classify_para(para):
        style_lower = para.style.name.lower() if para.style else ""
        for key, role in STYLE_ROLE.items():
            if key in style_lower:
                return role
        size   = _size_pt(para.runs)
        bold   = _all_bold(para.runs)
        italic = _all_italic(para.runs)
        if size and size in size_map:
            level = size_map[size]
            if level in ("title", "h1", "h2", "h3"):
                return level
        if bold: return "h3"
        text_lower = para.text.lower()
        if italic or any(kw in text_lower for kw in ("next:", "complete.", "coming up")):
            return "caption"
        return "body"

    def classify_table(tbl):
        nrows = len(tbl.rows)
        ncols = len(tbl.columns) if tbl.rows else 0
        first_text = first_size = None
        first_bold = False
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        first_text = t
                        first_size = _size_pt(p.runs)
                        first_bold = _all_bold(p.runs)
                        break
                if first_text: break
            if first_text: break
        if nrows > 2 and ncols >= 2: return "table"
        if ncols <= 1:
            if not first_bold: return "code"
            has_emoji  = first_text and any(ord(c) > 0x2000 for c in first_text[:3])
            size_level = size_map.get(first_size, "body") if first_size else "body"
            is_banner  = first_text and len(first_text) < 80 and not has_emoji and (
                first_text.upper() == first_text or first_text[:1].isdigit()
                or size_level in ("title", "h1", "h2")
            )
            return "h1" if is_banner else "body"
        return "table"

    def table_to_markdown(tbl) -> str:
        rows = []
        for row in tbl.rows:
            cells = [" ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()) for cell in row.cells]
            rows.append(cells)
        if not rows: return ""
        ncols = max(len(r) for r in rows)
        lines = ["| " + " | ".join(c.ljust(1) for c in rows[0]) + " |",
                 "| " + " | ".join(["---"] * ncols) + " |"]
        for row in rows[1:]:
            padded = row + [""] * (ncols - len(row))
            lines.append("| " + " | ".join(padded) + " |")
        return "\n".join(lines)

    def code_to_markdown(tbl) -> str:
        lines = []
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip(): lines.append(p.text)
        return "```python\n" + "\n".join(lines) + "\n```"

    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            para = para_map.get(child)
            if not para or not para.text.strip(): continue
            elements.append(Element(type=classify_para(para), text=para.text.strip(), page=0))
        elif tag == 'tbl':
            tbl = tbl_map.get(child)
            if not tbl: continue
            role = classify_table(tbl)
            if role == "code":
                elements.append(Element(type="code", text=code_to_markdown(tbl), page=0))
            elif role == "table":
                elements.append(Element(type="table", text=table_to_markdown(tbl), page=0))
            else:
                first_text = ""
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                first_text = p.text.strip()
                                break
                        if first_text: break
                    if first_text: break
                if first_text:
                    elements.append(Element(type=role, text=first_text, page=0))

    return elements


# ─────────────────────────────────────────────────────────────────────────────
# PDF ENGINE
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf(path: str) -> list:
    import fitz
    doc       = fitz.open(path)
    raw_paras = []

    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0: continue
            for line in block.get("lines", []):
                line_text = ""
                line_size = 0
                line_bold = False
                for span in line.get("spans", []):
                    t = span.get("text", "").strip()
                    if not t: continue
                    line_text += t + " "
                    line_size  = max(line_size, span.get("size", 0))
                    line_bold  = line_bold or bool(span.get("flags", 0) & 2**4)
                line_text = line_text.strip()
                if line_text:
                    raw_paras.append({"text": line_text, "size": round(line_size), "bold": line_bold, "page": page_num})

    if not raw_paras: return []

    size_counts    = Counter(p["size"] for p in raw_paras if p["size"])
    total          = len(raw_paras)
    body_threshold = max(total * 0.25, 3)
    body_sizes     = {sz for sz, cnt in size_counts.items() if cnt >= body_threshold}
    heading_sizes  = sorted([sz for sz in size_counts if sz not in body_sizes], reverse=True)
    size_map       = {sz: "body" for sz in body_sizes}
    for i, sz in enumerate(heading_sizes):
        size_map[sz] = ["title", "h1", "h2", "h3"][min(i, 3)]

    elements = []
    for p in raw_paras:
        level = size_map.get(p["size"], "body")
        if level in ("title", "h1", "h2", "h3"): role = level
        elif p["bold"]: role = "h3"
        else: role = "body"
        elements.append(Element(type=role, text=p["text"], page=p["page"]))

    return elements


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def parse_document(file_path: str) -> tuple:
    """
    Parse a PDF or DOCX file.
    Returns: (elements, doc_title, filetype)
    """
    fmt      = detect_format(file_path)
    elements = extract_docx(file_path) if fmt == "docx" else extract_pdf(file_path)
    title    = extract_title(elements)
    return elements, title, fmt
